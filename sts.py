"""давайте рассмотрим одну из классических баз данных, называемую "Ирисы Фишера". Она представляет собой набор данных о различных видах ирисов и включает в себя такие параметры как длина и ширина чашелистика и лепестка.

Вот две задачи, которые можно решить на основе этой базы данных и использования основных понятий статистики:

1. Определить средние значения и стандартные отклонения для длины и ширины чашелистика и лепестка для каждого вида ирисов (Setosa, Versicolor, Virginica).

2. Провести корреляционный анализ между параметрами длины и ширины чашелистика и лепестка для всех ирисов.

Для решения этих задач можно использовать вычисление средних значений, стандартных отклонений, корреляций и построение диаграмм рассеяния.

Надеюсь, это поможет вам начать практическое применение статистических понятий."""
import docx.shared
import pandas as pd
from docx import Document
import matplotlib.pyplot as plt

dataBase = pd.read_csv("iris.csv")
# sepal - чашелистник; petal - лепесток
""""sepal.length","sepal.width","petal.length","petal.width","variety" """
# Длинна чашелистника, ширина чашелистника, длинна лепестка, ширина лепестка, разновидность
unique_values = dataBase['variety'].unique()
meanDict = {}
std_dev = {}
correlation_dict = {}
for each in unique_values:
    # Вычисление среднего значения
    sepL = dataBase[dataBase["variety"] == f"{each}"]["sepal.length"].mean()
    sepW = dataBase[dataBase["variety"] == f"{each}"]["sepal.width"].mean()

    petL = dataBase[dataBase["variety"] == f"{each}"]["petal.length"].mean()
    petW = dataBase[dataBase["variety"] == f"{each}"]["petal.width"].mean()
    meanDict[each] = [sepL, sepW, petL, petW]  # фиксируем значения

    # вычисление стандартного отклонения
    stdsL = dataBase[dataBase["variety"] == f"{each}"]["sepal.length"].std()
    stdsW = dataBase[dataBase["variety"] == f"{each}"]["sepal.width"].std()

    stdpL = dataBase[dataBase["variety"] == f"{each}"]["petal.length"].std()
    stdpW = dataBase[dataBase["variety"] == f"{each}"]["petal.width"].std()
    std_dev[each] = [stdsL, stdsW, stdpL, stdpW]  # фиксируем значения

    # Вычисление корреляции между параметрами
    slpl = dataBase[dataBase["variety"] == f"{each}"]["sepal.length"].corr(dataBase["petal.length"])
    swpw = dataBase[dataBase["variety"] == f"{each}"]["sepal.width"].corr(dataBase["petal.width"])
    correlation_dict[each] = slpl, swpw

    # Построение диаграммы рассеяния
    fig, (ax1, ax2) = plt.subplots(1, 2)  # Создание графиков в строку
    ax1.scatter(dataBase[dataBase["variety"] == f"{each}"]["sepal.length"],
                dataBase[dataBase["variety"] == f"{each}"]["sepal.width"])  # Построение первой диаграммы рассеяния
    ax1.set_title(f'Диаграмма рассеяния {each}')
    ax1.set_xlabel('Длинна чашелистника')
    ax1.set_ylabel('Ширина чашелистника')
    plt.xlabel('sepal.length')
    plt.ylabel('petal.length')
    plt.savefig(f"{each} length")
    ax2.scatter(dataBase[dataBase["variety"] == f"{each}"]["petal.length"],
                dataBase[dataBase["variety"] == f"{each}"]["petal.width"])  # Построение второй диаграммы рассеяния
    ax2.set_title(f'Диаграмма рассеяния {each}')
    ax2.set_xlabel('Длинна лепестка')
    ax2.set_ylabel('Ширина лепестка')

    plt.savefig(f"{each} length")
print(meanDict)
print(std_dev)
print(correlation_dict)

# Создание нового документа
doc = Document()

# Добавление информации в документ
doc.add_heading('Работа с базой данных \"Ирисы Фишера', level=3)
doc.add_paragraph('Целью работы является освоение на практике основных статистических понятий.\n')
doc.add_paragraph("\tИрисы Фишера представляет собой набор данных о различных видах ирисов и включает в себя такие параметры как длина и ширина чашелистика и лепестка.\
Вот две задачи, которые можно решить на основе этой базы данных и использования основных понятий статистики:\
\n\t1. Определить средние значения и стандартные отклонения для длины и ширины чашелистика и лепестка для каждого вида ирисов (Setosa, Versicolor, Virginica).\
\n\t2. Провести корреляционный анализ между параметрами длины и ширины чашелистика и лепестка для всех ирисов.\
Для решения этих задач можно использовать вычисление средних значений, стандартных отклонений, корреляций и построение диаграмм рассеяния.\
.\n\n")
doc.add_paragraph(
    "1)\n\tЯ буду выполнять поставленную задачу используя язык програмирования python. "
    "Первым шагом импортирую библиотеки, которые мне понадобяться. "
    "Далее считываю файл с помощью библиотеки pandas. С помощью метода unique создадим список различных видов Ирисов."
    "Далее проходя по каждому уникальному названию выделяем из Базы дынных элементы с данным названием, фиксируем их среднее значения -отдельно по каждому из 4ех параметров\n\t")
img = doc.add_picture('exp0.png', width=docx.shared.Mm(110))
i = len(unique_values) - 1
while i >= 0:
    if i + 1 == len(unique_values):
        doc.add_paragraph(
            f"\nПо итогу получаем следующее следующие средние значния:\n {unique_values[i]} = {meanDict[unique_values[i]]}")
    else:
        doc.add_paragraph(f"{unique_values[i]} = {meanDict[unique_values[i]]}")
    i -= 1
doc.add_paragraph("длина чашелистика |ширина чашелистника|длина лепестка|ширина лепестка\n")
doc.add_paragraph("Теперь перейдем ко второй части первого пункта, задача заключается в поиске стандартного отклонения")
doc.add_paragraph("Делаем это следующим образом (в этом же цикле)")
img = doc.add_picture('exp1.png', width=docx.shared.Mm(110))
i = len(unique_values) - 1
while i >= 0:
    if i + 1 == len(unique_values):
        doc.add_paragraph("Получаем следующие значения:\n" + f"{unique_values[i]} = {std_dev[unique_values[i]]}")
    else:
        doc.add_paragraph(f"{unique_values[i]} = {std_dev[unique_values[i]]}")
    i -= 1
doc.add_paragraph("где 1ое значение в списке - стандартное отклонение для длины чашелистника")
doc.add_paragraph("где 2ое значение в списке - стандартное отклонение для ширины чашелистника")
doc.add_paragraph("где 3ие значение в списке - стандартное отклонение для длины лепестка")
doc.add_paragraph("где 4ое значение в списке - стандартное отклонение для ширины лепестка\n\n")
doc.add_paragraph(
    "2)Перейдем ко второй части задания, в которой найдем кореляцию между длиной и шириной чашелистников, и лепестков соответсвенно")
doc.add_paragraph("Как я это сделал можно найти на втором скриншоте")
doc.add_paragraph("Получаем следующие значения:")
i = len(unique_values) - 1
while i >= 0:
    if i + 1 == len(unique_values):
        doc.add_paragraph(
            "Получаем следующие значения:\n" + f"{unique_values[i]} = {correlation_dict[unique_values[i]]}")
    else:
        doc.add_paragraph(f"{unique_values[i]} = {std_dev[unique_values[i]]}")
    i -= 1
# Сохранение документа
doc.save('отчет.docx')
"""Диаграммы рассеивания (scatter plots) - это тип графика, который используется для визуализации взаимосвязи между двумя переменными. На диаграмме рассеивания каждая точка представляет наблюдение с определенными значениями по двум переменным, и их отображение позволяет оценить наличие или отсутствие связи между этими переменными.

Диаграммы рассеивания полезны для выявления паттернов и взаимосвязей в данных. Они также могут помочь выявить выбросы, группировки и кластеры в данных. Кроме того, они могут использоваться для оценки концентрации и распределения данных.

На диаграмме рассеивания ось x обычно представляет одну переменную, ось y - другую переменную. Точки на графике показывают, как изменяется значение переменной y при изменении значения переменной x, и насколько сильно переменные взаимосвязаны."""
